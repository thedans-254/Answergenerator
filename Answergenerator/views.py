from django.shortcuts import render
import openai
import os
from django.shortcuts import render
from django.http import FileResponse
from .forms import UploadFileForm
from PyPDF2 import PdfReader
from docx import Document
import pytesseract
from PIL import Image
from django.conf import settings
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

openai.api_key = settings.OPENAI_API_KEY


# Create your views here.
from django.shortcuts import render
from .forms import UploadFileForm

def upload_file(request):
    if request.method == 'POST':
        form = UploadFileForm(request.POST, request.FILES)
        if form.is_valid():
            uploaded_file = form.save()
            file_path = os.path.join(settings.MEDIA_ROOT, str(uploaded_file.file))
            file_type = uploaded_file.file.name.split(".")[-1].lower()

            text = extract_text_from_file(file_path, file_type)
            ai_response = get_answers_from_ai(text)

            print("AI RESPONSE:\n", ai_response)  # Debugging step

            # Attempt to split answers more reliably
            answers_list = ai_response.split("\n")  # Simple line-based split

            # Ensure we only keep non-empty responses
            answers_list = [answer.strip() for answer in answers_list if answer.strip()]

            # Preview first 5 answers (if available)
            preview_answers = answers_list[:10] if len(answers_list) >= 10 else answers_list

            # Generate output file
            output_filename = "exam_answers.pdf"
            output_file_path = save_as_pdf(ai_response, output_filename)

            return render(request, 'Answergenerator/results.html', {
                'preview_answers': preview_answers,
                'output_file': output_filename
            })

    else:
        form = UploadFileForm()
    return render(request, 'Answergenerator/upload.html', {'form': form})




# Extract text from different file formats
def extract_text_from_file(file_path, file_type):
    text = ""
    if file_type == "pdf":
        with open(file_path, "rb") as f:
            reader = PdfReader(f)
            for page in reader.pages:
                text += page.extract_text() + "\n"
    elif file_type == "docx":
        doc = Document(file_path)
        for para in doc.paragraphs:
            text += para.text + "\n"
    elif file_type in ["jpg", "jpeg", "png"]:
        text = pytesseract.image_to_string(Image.open(file_path))
    return text

# Generate AI answers using GPT-4 Turbo
def get_answers_from_ai(text):
    prompt = f"Answer the following exam questions from the text:\n\n{text}"
    response = openai.ChatCompletion.create(
        model="gpt-4-turbo",
        messages=[{"role": "system", "content": "You are a helpful AI that answers exam questions accurately."},
                  {"role": "user", "content": prompt}]
    )
    return response["choices"][0]["message"]["content"]

# Save answers as a PDF
def save_as_pdf(answers, filename):
    file_path = os.path.join(settings.MEDIA_ROOT, filename)
    c = canvas.Canvas(file_path, pagesize=letter)
    c.drawString(100, 750, "AI-Generated Answers")
    
    text_lines = answers.split("\n")
    y_position = 730
    for line in text_lines:
        c.drawString(100, y_position, line)
        y_position -= 15
        if y_position < 50:
            c.showPage()
            y_position = 750
    
    c.save()
    return file_path

# Save answers as a DOCX
def save_as_docx(answers, filename):
    file_path = os.path.join(settings.MEDIA_ROOT, filename)
    doc = Document()
    doc.add_heading("AI-Generated Answers", level=1)
    doc.add_paragraph(answers)
    doc.save(file_path)
    return file_path


# View to handle file download
def download_file(request, filename):
    file_path = os.path.join(settings.MEDIA_ROOT, filename)
    return FileResponse(open(file_path, 'rb'), as_attachment=True)
